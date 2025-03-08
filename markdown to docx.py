from docx import Document
from docx.shared import Inches
import os
import re

def clean_markdown_text(text):
    """Menghapus karakter markdown yang tidak perlu seperti **, _, dan spasi berlebih"""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Hapus bold **text**
    text = re.sub(r'_(.*?)_', r'\1', text)  # Hapus italic _text_
    text = re.sub(r'`(.*?)`', r'\1', text)  # Hapus inline code `text`
    return text.strip()

def convert_markdown_to_word(markdown_file, output_docx, image_folder="formula_images"):
    # Buka file Markdown
    with open(markdown_file, "r", encoding="utf-8") as md_file:
        markdown_content = md_file.readlines()

    # Buat dokumen Word baru
    doc = Document()

    # Regex untuk mendeteksi heading
    heading_pattern = re.compile(r'^(#{1,6})\s*(.*)')
    
    # Regex untuk mendeteksi daftar bernomor dan bullet
    numbered_list_pattern = re.compile(r'^\s*\d+\.\s+(.+)$')
    bullet_list_pattern = re.compile(r'^\s*[-*+]\s+(.+)$')
    
    # Regex untuk mendeteksi gambar rumus dalam format ![Rumus](path)
    formula_pattern = re.compile(r'!\[Rumus .*?\]\((.*?)\)')

    paragraph_buffer = []
    inside_code_block = False

    for line in markdown_content:
        line = line.rstrip()  # Hapus spasi atau newline ekstra di akhir baris

        # Deteksi awal atau akhir kode blok ```
        if line.startswith("```"):
            inside_code_block = not inside_code_block
            continue  # Lewati baris kode blok

        # Bersihkan teks jika bukan dalam kode blok
        if not inside_code_block:
            line = clean_markdown_text(line.strip())

        if not line:
            # Jika ada paragraf yang terkumpul, tambahkan ke dokumen
            if paragraph_buffer:
                doc.add_paragraph(" ".join(paragraph_buffer))
                paragraph_buffer = []
            continue
        
        # Deteksi heading
        heading_match = heading_pattern.match(line)
        if heading_match:
            if paragraph_buffer:
                doc.add_paragraph(" ".join(paragraph_buffer))
                paragraph_buffer = []
            heading_level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            doc.add_heading(heading_text, level=heading_level)
            continue
        
        # Deteksi daftar bernomor
        if numbered_list_pattern.match(line):
            if paragraph_buffer:
                doc.add_paragraph(" ".join(paragraph_buffer))
                paragraph_buffer = []
            doc.add_paragraph(line, style='List Number')
            continue
        
        # Deteksi daftar bullet
        if bullet_list_pattern.match(line):
            if paragraph_buffer:
                doc.add_paragraph(" ".join(paragraph_buffer))
                paragraph_buffer = []
            doc.add_paragraph(line, style='List Bullet')
            continue

        # Deteksi gambar rumus
        formula_match = formula_pattern.findall(line)
        if formula_match:
            if paragraph_buffer:
                doc.add_paragraph(" ".join(paragraph_buffer))
                paragraph_buffer = []

            # Proses teks di sekitar gambar
            text_parts = re.split(formula_pattern, line)
            for idx, part in enumerate(text_parts):
                if idx % 2 == 0:
                    if part.strip():
                        doc.add_paragraph(part.strip())
                else:
                    image_name = os.path.basename(part.strip())  # Ambil nama gambar dari Markdown
                    image_path = os.path.join(image_folder, image_name)

                    if os.path.exists(image_path):
                        doc.add_paragraph(f"[Gambar: {image_name}]")
                        doc.add_picture(image_path, width=Inches(2))
                    else:
                        doc.add_paragraph(f"[Gambar tidak ditemukan: {image_name}]")
                        print(f"⚠️ Gambar tidak ditemukan: {image_path}")

            continue
        
        # Jika bukan bagian khusus (heading, list, gambar), tambahkan ke buffer paragraf
        paragraph_buffer.append(line)

    # Tambahkan paragraf terakhir jika masih ada dalam buffer
    if paragraph_buffer:
        doc.add_paragraph(" ".join(paragraph_buffer))
    
    # Simpan dokumen Word
    doc.save(output_docx)
    print(f"✅ File Word berhasil disimpan di: {output_docx}")

# Contoh penggunaan
markdown_file = "output.md"  # Ganti dengan path file Markdown Anda
output_docx = "document_output.docx"  # File Word output
convert_markdown_to_word(markdown_file, output_docx, image_folder="formula_images")
